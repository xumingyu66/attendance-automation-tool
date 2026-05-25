"""Typed wrapper for python-docx — generates Word attendance reports.

Matches the formatting of the reference template:
  - Title: 图书馆在岗学生馆员YYYY年M月工作时长统计表 (17pt)
  - Subtitle: (值班岗) / (整架岗) (16pt)
  - Tables: 宋体 14pt, 6-col for 值班岗, 3-col for 整架岗
"""

from __future__ import annotations

import calendar
import datetime

from docx import Document
from docx.shared import Pt, Cm, Emu
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn


class DocxReportGenerator:
    """Generates a formatted Word attendance report matching the reference template."""

    # Font sizes (in Pt)
    TITLE_SIZE = Pt(17)
    SUBTITLE_SIZE = Pt(16)
    TABLE_SIZE = Pt(14)

    FONT_NAME = "宋体"

    def generate(
        self,
        output_path: str,
        year: int,
        month: int,
        zhengjia_data: list[dict],
        zhiban_data: list[dict],
    ) -> None:

        doc = Document()

        # Page setup — match template
        section = doc.sections[0]
        section.page_width = Emu(7560310)
        section.page_height = Emu(10692130)
        section.left_margin = Emu(1143000)
        section.right_margin = Emu(1143000)
        section.top_margin = Emu(914400)
        section.bottom_margin = Emu(914400)

        # ---- 值班岗 section ----
        self._add_title(doc, year, month)
        self._add_subtitle(doc, "（值班岗）")
        self._write_table_6col(doc, zhiban_data)

        doc.add_paragraph()

        # ---- 整架岗 section ----
        self._add_title(doc, year, month)
        self._add_subtitle(doc, "（整架岗）")
        self._write_table_3col(doc, zhengjia_data)

        # ---- Footer ----
        doc.add_paragraph()
        self._add_footer(doc)

        doc.save(output_path)

    # =========================================================================
    # Title / Subtitle / Footer
    # =========================================================================

    def _add_title(self, doc: Document, year: int, month: int) -> None:
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run(f"图书馆在岗学生馆员{year}年{month}月工作时长统计表")
        run.font.size = self.TITLE_SIZE
        self._set_font(run)

    def _add_subtitle(self, doc: Document, text: str) -> None:
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = self.SUBTITLE_SIZE
        self._set_font(run)

    def _add_footer(self, doc: Document) -> None:
        today = datetime.date.today()
        p = doc.add_paragraph()
        run = p.add_run(f"填报人：          {today.year}年{today.month}月{today.day}日")
        run.font.size = self.TABLE_SIZE
        self._set_font(run)

    # =========================================================================
    # Table: 值班岗 — 6 columns (2 employees per row)
    # =========================================================================

    def _write_table_6col(self, doc: Document, data: list[dict]) -> None:
        """6-column table: 序号|姓名|工作时长(小时)|序号|姓名|工作时长(小时)

        Employees are paired left-right per row.
        """
        table = doc.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        # Header
        hdrs = ["序号", "姓名", "工作时长\n（小时）", "序号", "姓名", "工作时长\n（小时）"]
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(hdrs):
            self._set_cell(hdr_cells[i], h, bold=True)

        # Pair employees: left in columns 0-2, right in columns 3-5
        for idx in range(0, len(data), 2):
            left = data[idx]
            right = data[idx + 1] if idx + 1 < len(data) else None

            row = table.add_row()
            cells = row.cells

            # Left side
            self._set_cell(cells[0], str(idx + 1))
            self._set_cell(cells[1], left["name"])
            self._set_cell(cells[2], str(left["total_hours"]))

            # Right side (序号 continues from left)
            if right:
                self._set_cell(cells[3], str(idx + 2))
                self._set_cell(cells[4], right["name"])
                self._set_cell(cells[5], str(right["total_hours"]))
            else:
                self._set_cell(cells[3], "")
                self._set_cell(cells[4], "")
                self._set_cell(cells[5], "")

    # =========================================================================
    # Table: 整架岗 — 3 columns (1 employee per row)
    # =========================================================================

    def _write_table_3col(self, doc: Document, data: list[dict]) -> None:
        """3-column table: 序号|姓名|工作时长（小时）"""
        table = doc.add_table(rows=1, cols=3)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER

        hdrs = ["序号", "姓名", "工作时长（小时）"]
        hdr_cells = table.rows[0].cells
        for i, h in enumerate(hdrs):
            self._set_cell(hdr_cells[i], h, bold=True)

        for idx, d in enumerate(data):
            row = table.add_row()
            cells = row.cells
            self._set_cell(cells[0], str(idx + 1))
            self._set_cell(cells[1], d["name"])
            self._set_cell(cells[2], str(d["total_hours"]))

    # =========================================================================
    # Cell formatting helpers
    # =========================================================================

    def _set_cell(self, cell, text: str, bold: bool = False) -> None:
        cell.text = ""
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(text)
        run.font.size = self.TABLE_SIZE
        run.bold = bold
        self._set_font(run)

    def _set_font(self, run) -> None:
        run.font.name = self.FONT_NAME
        run._element.rPr.rFonts.set(qn("w:eastAsia"), self.FONT_NAME)
